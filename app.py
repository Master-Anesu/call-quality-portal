"""
Call Quality Review Portal — Trilogy Care
Flask app that automates call quality reviews for sales reps.
Uses background jobs + polling instead of SSE to avoid Render request timeouts.
"""

import os
import sys
import json
import re
import uuid
import base64
import logging
import threading
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor

from flask import Flask, render_template, jsonify, request, send_from_directory

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '.brain', 'common_functions'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '.brain', 'templates', 'branded', 'base'))

logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['OUTPUT_DIR'] = os.path.join(os.path.dirname(__file__), 'output')

# ---------------------------------------------------------------------------
# In-memory job store  (single-worker gunicorn, so this is safe)
# ---------------------------------------------------------------------------
jobs = {}

# ---------------------------------------------------------------------------
# Config from env vars
# ---------------------------------------------------------------------------
MCP_BASE_URL = "https://llm-alb.trilogycare.com.au/mcp/tools"
MCP_API_KEY = os.environ.get("MCP_API_KEY", "")
WORKSPACE_USER_ID = os.environ.get("WORKSPACE_USER_ID", "")
SENDER_EMAIL = "anesut@trilogycare.com.au"

AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY", "")
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT", "")
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2025-01-01-preview")
AZURE_OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT_NAME", "trilogy-gpt-4.1")

AZURE_AD_TENANT_ID = os.environ.get("AZURE_AD_TENANT_ID", "")
AZURE_AD_CLIENT_ID = os.environ.get("AZURE_AD_CLIENT_ID", "")
AZURE_AD_CLIENT_SECRET = os.environ.get("AZURE_AD_CLIENT_SECRET", "")

DATABRICKS_HOST = os.environ.get("DATABRICKS_HOST", "")
DATABRICKS_TOKEN = os.environ.get("DATABRICKS_TOKEN", "")
DATABRICKS_WAREHOUSE_ID = os.environ.get("DATABRICKS_WAREHOUSE_ID", "")


# ---------------------------------------------------------------------------
# MCP Response Parsers
# ---------------------------------------------------------------------------

def _parse_aircall_calls(result) -> list:
    """Extract calls list from various MCP response shapes."""
    if not result or (isinstance(result, dict) and result.get('error')):
        return []
    if isinstance(result, list):
        return result
    if isinstance(result, dict):
        if 'calls' in result and isinstance(result['calls'], list):
            return result['calls']
        if 'result' in result:
            inner = result['result']
            if isinstance(inner, list):
                return inner
            if isinstance(inner, dict) and 'calls' in inner:
                return inner['calls']
        if 'content' in result:
            content = result['content']
            if isinstance(content, list):
                for block in content:
                    if isinstance(block, dict) and block.get('type') == 'text':
                        try:
                            parsed = json.loads(block['text'])
                            if isinstance(parsed, list):
                                return parsed
                            if isinstance(parsed, dict):
                                if 'calls' in parsed:
                                    return parsed['calls']
                                if 'result' in parsed and isinstance(parsed['result'], dict):
                                    return parsed['result'].get('calls', [])
                        except (json.JSONDecodeError, TypeError):
                            pass
            elif isinstance(content, str):
                try:
                    parsed = json.loads(content)
                    if isinstance(parsed, list):
                        return parsed
                    if isinstance(parsed, dict) and 'calls' in parsed:
                        return parsed['calls']
                except (json.JSONDecodeError, TypeError):
                    pass
        if 'raw' in result and isinstance(result['raw'], str):
            try:
                parsed = json.loads(result['raw'])
                if isinstance(parsed, list):
                    return parsed
                if isinstance(parsed, dict) and 'calls' in parsed:
                    return parsed['calls']
            except (json.JSONDecodeError, TypeError):
                pass
    return []


# ---------------------------------------------------------------------------
# MCP Tool Caller
# ---------------------------------------------------------------------------

def call_mcp_tool(tool_name: str, arguments: dict) -> dict:
    """Call an MCP tool via HTTP."""
    import requests as req_lib
    url = f"{MCP_BASE_URL}/{tool_name}/invoke"
    headers = {
        "X-API-Key": MCP_API_KEY,
        "X-User-ID": WORKSPACE_USER_ID,
        "Content-Type": "application/json",
    }
    try:
        resp = req_lib.post(url, json={"arguments": arguments}, headers=headers, timeout=30)
        resp.raise_for_status()
        if not resp.text.strip():
            return {"success": True}
        try:
            return resp.json()
        except ValueError:
            # HTTP 2xx but non-JSON body — the operation succeeded
            return {"success": True, "raw": resp.text[:500]}
    except Exception as e:
        return {"error": str(e)}


# ---------------------------------------------------------------------------
# Databricks direct query (for transcripts not yet in MCP cache)
# ---------------------------------------------------------------------------

def query_databricks(sql: str) -> list:
    """Execute a SQL query against Databricks and return rows."""
    import requests as req_lib
    if not all([DATABRICKS_HOST, DATABRICKS_TOKEN, DATABRICKS_WAREHOUSE_ID]):
        return []
    try:
        resp = req_lib.post(
            f"{DATABRICKS_HOST}/api/2.0/sql/statements/",
            headers={"Authorization": f"Bearer {DATABRICKS_TOKEN}", "Content-Type": "application/json"},
            json={"warehouse_id": DATABRICKS_WAREHOUSE_ID, "statement": sql, "wait_timeout": "30s"},
            timeout=35,
        )
        data = resp.json()
        if data.get('status', {}).get('state') == 'SUCCEEDED':
            return data.get('result', {}).get('data_array', [])
    except Exception as e:
        logger.error("Databricks query failed: %s", e)
    return []


def normalize_phone(phone: str) -> str:
    """Normalize a phone number to digits only (last 9 digits) for matching."""
    if not phone or phone == 'None' or phone == 'null':
        return ''
    digits = re.sub(r'\D', '', phone)
    # Use last 9 digits to handle country code variations (+61 vs 0)
    return digits[-9:] if len(digits) >= 9 else digits


def get_active_lead_phones() -> set:
    """Get normalized phone numbers for leads with Active or Allocated journey stage."""
    rows = query_databricks("""
        SELECT Phone, Mobile, Contact_Phone
        FROM trilogycare_dev.zoho_crm.leads
        WHERE Journey_Stage IN ('B-Active HCP', 'B-Allocated HCP')
    """)
    phones = set()
    for row in rows:
        for val in row:
            if val:
                n = normalize_phone(str(val))
                if n:
                    phones.add(n)
    return phones


def get_deal_phones_by_date() -> dict:
    """Get mapping of normalized phone -> set of deal creation dates (YYYY-MM-DD).
    Only includes deals from the last 90 days for performance."""
    rows = query_databricks("""
        SELECT Client_Phone, Mobile, CAST(Created_Time AS STRING)
        FROM trilogycare_dev.zoho_crm.deals
        WHERE Created_Time >= CURRENT_DATE - INTERVAL 90 DAYS
        AND (Client_Phone IS NOT NULL OR Mobile IS NOT NULL)
    """)
    phone_dates = {}
    for row in rows:
        created_str = row[2] or ''
        deal_date = created_str[:10] if created_str else ''
        if not deal_date:
            continue
        for val in row[:2]:
            if val:
                n = normalize_phone(str(val))
                if n:
                    if n not in phone_dates:
                        phone_dates[n] = set()
                    phone_dates[n].add(deal_date)
    return phone_dates


def is_call_eligible(caller_phone: str, call_date: str, lead_phones: set, deal_phone_dates: dict) -> bool:
    """Check if a call is eligible for review.
    Eligible if:
    - Caller is an Active/Allocated lead, OR
    - A deal was created for this phone on the same day as the call
    """
    normalized = normalize_phone(caller_phone)
    if not normalized:
        return False
    # Active/Allocated lead — always eligible
    if normalized in lead_phones:
        return True
    # Deal created on the same day as the call
    if normalized in deal_phone_dates and call_date in deal_phone_dates[normalized]:
        return True
    return False


def check_call_eligible(phone: str, call_date: str = '') -> dict:
    """Check if a specific call is eligible for review.
    Used by the review pipeline for individual call verification."""
    if not phone:
        return {'is_valid': False, 'source': None, 'lead_name': None}
    normalized = normalize_phone(phone)
    if not normalized:
        return {'is_valid': False, 'source': None, 'lead_name': None}

    # Check leads (Active/Allocated)
    rows = query_databricks(f"""
        SELECT First_Name, Last_Name, Journey_Stage
        FROM trilogycare_dev.zoho_crm.leads
        WHERE Journey_Stage IN ('B-Active HCP', 'B-Allocated HCP')
        AND (
            RIGHT(REGEXP_REPLACE(Phone, '[^0-9]', ''), 9) = '{normalized}'
            OR RIGHT(REGEXP_REPLACE(Mobile, '[^0-9]', ''), 9) = '{normalized}'
            OR RIGHT(REGEXP_REPLACE(Contact_Phone, '[^0-9]', ''), 9) = '{normalized}'
        )
        LIMIT 1
    """)
    if rows:
        return {
            'is_valid': True,
            'source': 'lead',
            'journey_stage': rows[0][2] or '',
            'lead_name': f"{rows[0][0] or ''} {rows[0][1] or ''}".strip(),
        }

    # Check deals created on the same day as the call
    if call_date:
        rows = query_databricks(f"""
            SELECT Contact_Name.name, Stage, CAST(Created_Time AS STRING)
            FROM trilogycare_dev.zoho_crm.deals
            WHERE CAST(Created_Time AS DATE) = '{call_date}'
            AND (
                RIGHT(REGEXP_REPLACE(Client_Phone, '[^0-9]', ''), 9) = '{normalized}'
                OR RIGHT(REGEXP_REPLACE(Mobile, '[^0-9]', ''), 9) = '{normalized}'
            )
            LIMIT 1
        """)
        if rows:
            return {
                'is_valid': True,
                'source': 'deal',
                'journey_stage': f"Deal - {rows[0][1] or ''}",
                'lead_name': rows[0][0] or '',
            }

    return {'is_valid': False, 'source': None, 'lead_name': None}


def get_transcript_from_databricks(call_id: str) -> str:
    """Fetch transcript directly from Databricks transcriptions table."""
    rows = query_databricks(f"""
        SELECT content.utterances
        FROM trilogycare_dev.aircall.transcriptions
        WHERE call_id = {call_id}
        LIMIT 1
    """)
    if not rows or not rows[0][0]:
        return ''
    utterances = rows[0][0]
    if isinstance(utterances, str):
        try:
            utterances = json.loads(utterances)
        except (json.JSONDecodeError, TypeError):
            return ''
    if isinstance(utterances, list):
        return '\n'.join(
            f"[{u.get('participant_type', 'unknown')}]: {u.get('text', '')}"
            for u in utterances if u.get('text')
        )
    return ''


def cache_recording(call_id: str, recording_url: str) -> str:
    """Download recording to local cache while URL is still valid. Returns local path or empty."""
    import requests as req_lib
    cache_dir = os.path.join(app.config['OUTPUT_DIR'], 'recordings')
    os.makedirs(cache_dir, exist_ok=True)
    cache_path = os.path.join(cache_dir, f"{call_id}.mp3")
    if os.path.isfile(cache_path):
        return cache_path
    try:
        resp = req_lib.get(recording_url, timeout=120)
        if resp.status_code == 200 and len(resp.content) > 1000:
            with open(cache_path, 'wb') as f:
                f.write(resp.content)
            logger.info("Cached recording for call %s (%d bytes)", call_id, len(resp.content))
            return cache_path
    except Exception as e:
        logger.warning("Failed to cache recording for call %s: %s", call_id, e)
    return ''


# ---------------------------------------------------------------------------
# Recording transcription (fallback when no stored transcript exists)
# ---------------------------------------------------------------------------

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")


def get_call_details(call_id: str) -> dict:
    """Fetch full call details including transcript segments and recording URL.
    Uses get_aircall_call_details which returns everything in one request."""
    result = call_mcp_tool('get_aircall_call_details', {'call_id': str(call_id)})
    if isinstance(result, dict):
        # Unwrap MCP content envelope
        if 'content' in result:
            content = result['content']
            if isinstance(content, list):
                for block in content:
                    if isinstance(block, dict) and block.get('type') == 'text':
                        try:
                            result = json.loads(block['text'])
                            break
                        except (json.JSONDecodeError, TypeError):
                            pass
        inner = result.get('result', result)
        if isinstance(inner, dict):
            return inner.get('call', inner)
    return {}


def get_fresh_recording_url(call_id: str) -> str:
    """Get a recording URL for a call. Note: MCP tools return pre-signed S3 URLs
    that may be expired. This fetches whatever is available."""
    call_data = get_call_details(call_id)
    return call_data.get('recording', '') or call_data.get('recording_url', '') or ''


def transcribe_recording(recording_url: str, call_id: str = '') -> str:
    """Transcribe a call recording via Groq Whisper. Checks local cache first."""
    import requests as req_lib
    import tempfile
    if not GROQ_API_KEY:
        logger.warning("No GROQ_API_KEY configured — cannot transcribe recording")
        return ""

    tmp_path = None
    try:
        # Check for cached local recording first
        if call_id:
            cache_dir = os.path.join(app.config['OUTPUT_DIR'], 'recordings')
            cached = os.path.join(cache_dir, f"{call_id}.mp3")
            if os.path.isfile(cached):
                logger.info("Using cached recording for call %s", call_id)
                tmp_path = cached

        # If no URL provided but we have a call_id, fetch a fresh URL
        if not tmp_path and not recording_url and call_id:
            logger.info("No recording URL provided for call %s, fetching fresh URL...", call_id)
            recording_url = get_fresh_recording_url(call_id)

        # Handle local file path
        if not tmp_path and recording_url and os.path.isfile(recording_url):
            tmp_path = recording_url

        # Download from URL
        if not tmp_path and recording_url:
            resp = req_lib.get(recording_url, timeout=120)
            if resp.status_code in (403, 400) and call_id:
                logger.info('Recording URL expired for call %s, fetching fresh URL...', call_id)
                fresh_url = get_fresh_recording_url(call_id)
                if fresh_url:
                    recording_url = fresh_url
                    resp = req_lib.get(recording_url, timeout=120)
                else:
                    logger.error('Could not get fresh recording URL for call %s', call_id)
                    return ""
            resp.raise_for_status()
            with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp:
                tmp.write(resp.content)
                tmp_path = tmp.name

        if not tmp_path:
            return ""

        with open(tmp_path, "rb") as audio_file:
            result = req_lib.post(
                "https://api.groq.com/openai/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                files={"file": ("recording.mp3", audio_file, "audio/mpeg")},
                data={"model": "whisper-large-v3", "language": "en", "response_format": "text"},
                timeout=300,
            )
            result.raise_for_status()

        # Don't delete cached recordings
        cache_dir = os.path.join(app.config['OUTPUT_DIR'], 'recordings')
        if tmp_path and not tmp_path.startswith(cache_dir) and tmp_path != recording_url:
            os.unlink(tmp_path)

        return result.text.strip() or ""
    except Exception as e:
        logger.error("Failed to transcribe recording: %s", e)
        return ""


# ---------------------------------------------------------------------------
# Azure OpenAI  (non-streaming — background thread handles the wait)
# ---------------------------------------------------------------------------

def call_llm(system_prompt: str, user_prompt: str, max_tokens: int = 4000) -> str:
    """Call Azure OpenAI GPT-4.1 (non-streaming). Returns the full response text."""
    from openai import AzureOpenAI
    client = AzureOpenAI(
        api_key=AZURE_OPENAI_API_KEY,
        api_version=AZURE_OPENAI_API_VERSION,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
    )
    response = client.chat.completions.create(
        model=AZURE_OPENAI_DEPLOYMENT,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_tokens=max_tokens,
        temperature=0.3,
        stream=False,
        response_format={"type": "json_object"},
    )
    return response.choices[0].message.content or ""


# ---------------------------------------------------------------------------
# Scoring prompt
# ---------------------------------------------------------------------------

SCORING_SYSTEM_PROMPT = """You are a call quality reviewer for Trilogy Care, an Australian aged care provider.
You score sales/onboarding calls across 4 stages of a sales call. This is a coaching tool, not a performance audit.
Be specific — ground every score and comment in a real moment from the transcript.

IMPORTANT: All feedback must speak DIRECTLY to the rep using "you" / "your" (second person). Never refer to them in third person ("the rep", "she", "he", "they did"). Write as if you are coaching them face-to-face. Example: "You built rapport quickly by using their name" NOT "The rep built rapport by using their name".

Scoring benchmark: top-performer Aprocina Anthony's conversational style — the way she builds trust, reads callers, and makes progress feel natural.
Core question for every stage: "Did this part of the conversation move the caller closer to a confident yes?"

The call is scored across 4 stages with these weightings:
- Introduction (15 points / 15%): Sets the tone
- Discovery (35 points / 35%): The foundation — everything else depends on this
- Pitch (35 points / 35%): Where value is communicated and trust is built
- Close (15 points / 15%): Seals what was already earned

Return your response as valid JSON with this exact structure:
{
  "client_name": "the client/caller's full name as identified from the transcript",
  "scores": {
    "introduction": {
      "total": 0,
      "sub_scores": {
        "warm_greeting": {"score": 0, "max": 5, "comment": "specific feedback"},
        "clear_introduction": {"score": 0, "max": 5, "comment": "specific feedback"},
        "purpose_statement": {"score": 0, "max": 5, "comment": "specific feedback"}
      },
      "feedback": "2-3 sentence paragraph speaking directly to the rep (use 'you') — what you did well or missed, specific moments, how it affected discovery"
    },
    "discovery": {
      "total": 0,
      "sub_scores": {
        "current_care_situation": {"score": 0, "max": 8, "comment": "specific feedback"},
        "biggest_challenge": {"score": 0, "max": 8, "comment": "specific feedback"},
        "impact_on_daily_life": {"score": 0, "max": 7, "comment": "specific feedback"},
        "ideal_care_vision": {"score": 0, "max": 6, "comment": "specific feedback"},
        "timeline_urgency": {"score": 0, "max": 6, "comment": "specific feedback"}
      },
      "triggers_identified": ["list of triggers the rep uncovered, or note what was missed"],
      "feedback": "3-4 sentence paragraph speaking directly to the rep (use 'you') — did you uncover 3 clear triggers? How did this affect the pitch?"
    },
    "pitch": {
      "total": 0,
      "sub_scores": {
        "trigger_recap": {"score": 0, "max": 8, "comment": "specific feedback"},
        "solution_alignment": {"score": 0, "max": 10, "comment": "specific feedback"},
        "relevant_services": {"score": 0, "max": 9, "comment": "specific feedback"},
        "comprehension_check": {"score": 0, "max": 8, "comment": "specific feedback"}
      },
      "feedback": "3-4 sentence paragraph speaking directly to the rep (use 'you') — was your pitch built on triggers or generic? Tie back to discovery."
    },
    "close": {
      "total": 0,
      "sub_scores": {
        "closing_question": {"score": 0, "max": 4, "comment": "specific feedback"},
        "pause_and_listen": {"score": 0, "max": 3, "comment": "specific feedback"},
        "objection_handling": {"score": 0, "max": 4, "comment": "specific feedback"},
        "next_steps_confirmed": {"score": 0, "max": 4, "comment": "specific feedback"}
      },
      "closing_technique": "process / pricing / alternate / assumptive",
      "feedback": "2-3 sentence paragraph speaking directly to the rep (use 'you') — was your close confident or hesitant? Trace back to earlier stages if needed."
    }
  },
  "stage_flow": "3-4 sentence paragraph speaking directly to the rep (use 'you') — connecting the dots across all four stages, where your momentum built or broke down",
  "focus_on": "one sentence speaking directly to the rep — the single most important thing for you to work on next",
  "top_strength": "one sentence speaking directly to the rep — the one thing you should keep doing",
  "top_development_area": "one sentence speaking directly to the rep — the one thing that would make the biggest difference for you"
}

Total max: 100 points (Introduction /15 + Discovery /35 + Pitch /35 + Close /15)

Grade scale:
A+ = 90-100%, A = 80-89%, B = 70-79%, C = 60-69%, D = below 60%
"""


# ---------------------------------------------------------------------------
# Word Document Generator
# ---------------------------------------------------------------------------

def generate_word_doc(review_data: dict) -> str:
    """Generate a branded TC Word document from review data. Returns file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    NAVY = RGBColor(0x2C, 0x4C, 0x79)
    TEAL = RGBColor(0x43, 0xC0, 0xBE)
    DARK_TEAL = RGBColor(0x00, 0x7F, 0x7E)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLACK = RGBColor(0x33, 0x33, 0x33)

    doc = Document()

    today_str = datetime.now().strftime('%d %B %Y')

    # Header block (no cover page per spec)
    header_data = [
        ['Agent', f"{review_data['rep_name']} — {review_data.get('est_id', '')} {review_data['rep_role']}"],
        ['Department', 'CSR'],
        ['Manager', 'Anesu Taderera'],
        ['Review Date', today_str],
        ['Call Reviewed', f"{review_data.get('call_date', '')} — {review_data.get('client_name', 'Unknown')} ({review_data.get('direction', '')}, {review_data.get('duration', '')} min)"],
    ]

    table = doc.add_table(rows=len(header_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(header_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = ''
        cell_value.text = ''
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = NAVY
        run_l.font.name = 'Aptos'
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.name = 'Aptos'
        run_v.font.color.rgb = BLACK
        shading = cell_label._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'E8EDF3',
        })
        shading.append(shading_elem)

    doc.add_paragraph('')

    # Calculate totals
    scores = review_data.get('scores', {})
    intro_total = scores.get('introduction', {}).get('total', 0)
    discovery_total = scores.get('discovery', {}).get('total', 0)
    pitch_total = scores.get('pitch', {}).get('total', 0)
    close_total = scores.get('close', {}).get('total', 0)
    total_points = intro_total + discovery_total + pitch_total + close_total
    percentage = round((total_points / 100) * 100, 1) if total_points else 0

    if percentage >= 90:
        grade = 'A+'
    elif percentage >= 80:
        grade = 'A'
    elif percentage >= 70:
        grade = 'B'
    elif percentage >= 60:
        grade = 'C'
    else:
        grade = 'D'

    # Section 1: Stage-by-Stage Feedback
    h1 = doc.add_heading('Stage-by-Stage Feedback', level=1)
    for run in h1.runs:
        run.font.color.rgb = NAVY

    # Introduction
    h2 = doc.add_heading(f"Introduction ({intro_total}/15)", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY
    intro_feedback = scores.get('introduction', {}).get('feedback', '')
    if intro_feedback:
        p = doc.add_paragraph(intro_feedback)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(11)

    # Discovery
    h2 = doc.add_heading(f"Discovery ({discovery_total}/35)", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY
    disc_feedback = scores.get('discovery', {}).get('feedback', '')
    if disc_feedback:
        p = doc.add_paragraph(disc_feedback)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
    triggers = scores.get('discovery', {}).get('triggers_identified', [])
    if triggers:
        p = doc.add_paragraph('')
        run_label = p.add_run('Triggers identified: ')
        run_label.bold = True
        run_label.font.name = 'Aptos'
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = NAVY
        run_text = p.add_run(', '.join(triggers) if isinstance(triggers, list) else str(triggers))
        run_text.font.name = 'Aptos'
        run_text.font.size = Pt(11)

    # Pitch
    h2 = doc.add_heading(f"Pitch ({pitch_total}/35)", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY
    pitch_feedback = scores.get('pitch', {}).get('feedback', '')
    if pitch_feedback:
        p = doc.add_paragraph(pitch_feedback)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(11)

    # Close
    h2 = doc.add_heading(f"Close ({close_total}/15)", level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY
    close_feedback = scores.get('close', {}).get('feedback', '')
    if close_feedback:
        p = doc.add_paragraph(close_feedback)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
    close_technique = scores.get('close', {}).get('closing_technique', '')
    if close_technique:
        p = doc.add_paragraph('')
        run_label = p.add_run('Closing technique: ')
        run_label.bold = True
        run_label.font.name = 'Aptos'
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = NAVY
        run_text = p.add_run(close_technique)
        run_text.font.name = 'Aptos'
        run_text.font.size = Pt(11)

    # Stage Flow
    h2 = doc.add_heading('Stage Flow', level=2)
    for run in h2.runs:
        run.font.color.rgb = NAVY
    stage_flow = review_data.get('stage_flow', '')
    if stage_flow:
        p = doc.add_paragraph(stage_flow)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(11)

    doc.add_paragraph('')

    # Section 2: Score Card
    h1 = doc.add_heading('Score Card', level=1)
    for run in h1.runs:
        run.font.color.rgb = NAVY

    stage_rows = [
        ('Introduction', '15%', f"{intro_total}/15"),
        ('Discovery', '35%', f"{discovery_total}/35"),
        ('Pitch', '35%', f"{pitch_total}/35"),
        ('Close', '15%', f"{close_total}/15"),
    ]

    score_table = doc.add_table(rows=len(stage_rows) + 2, cols=3)
    score_table.style = 'Table Grid'
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Stage', 'Weight', 'Score']
    for j, h_text in enumerate(headers):
        cell = score_table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Aptos'
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '2C4C79',
        })
        shading.append(shading_elem)

    for i, (stage_name, weight, score_str) in enumerate(stage_rows, start=1):
        row = score_table.row_cells(i)
        values = [stage_name, weight, score_str]
        for j, val in enumerate(values):
            row[j].text = ''
            run = row[j].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Aptos'
            if i % 2 == 0:
                shading = row[j]._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F5F5',
                })
                shading.append(shading_elem)

    # Total row
    total_row = score_table.row_cells(len(stage_rows) + 1)
    total_values = ['Total', '100%', f"{total_points}/100 — {grade}"]
    for j, val in enumerate(total_values):
        total_row[j].text = ''
        run = total_row[j].paragraphs[0].add_run(val)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Aptos'
        run.font.color.rgb = NAVY
        shading = total_row[j]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D4E6E6',
        })
        shading.append(shading_elem)

    doc.add_paragraph('')

    # Section 3: Focus On
    h1 = doc.add_heading('Focus On', level=1)
    for run in h1.runs:
        run.font.color.rgb = NAVY
    focus_on = review_data.get('focus_on', '')
    if focus_on:
        p = doc.add_paragraph(focus_on)
        for run in p.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(12)
            run.bold = True

    # Footer
    doc.add_paragraph('')
    doc.add_paragraph('_' * 60)
    footer_items = [
        ('Prepared by:', 'Trilogy Care Assistant (AI-generated review)'),
        ('Reviewed by:', 'Anesu Taderera'),
        ('Date:', today_str),
    ]
    for label, value in footer_items:
        p = doc.add_paragraph('')
        run_l = p.add_run(f"{label} ")
        run_l.bold = True
        run_l.font.size = Pt(9)
        run_l.font.name = 'Aptos'
        run_l.font.color.rgb = NAVY
        run_v = p.add_run(value)
        run_v.font.size = Pt(9)
        run_v.font.name = 'Aptos'

    # Save
    first = review_data['rep_name'].split()[0] if review_data['rep_name'] else 'Unknown'
    last = review_data['rep_name'].split()[-1] if len(review_data['rep_name'].split()) > 1 else ''
    month_year = datetime.now().strftime('%b%Y')
    filename = f"{first}_{last}_Call_Quality_Review_{month_year}.docx"
    filepath = os.path.join(app.config['OUTPUT_DIR'], filename)
    os.makedirs(app.config['OUTPUT_DIR'], exist_ok=True)
    doc.save(filepath)

    review_data['filename'] = filename
    review_data['filepath'] = filepath
    review_data['grade'] = grade
    review_data['total_points'] = total_points
    review_data['percentage'] = percentage

    return filepath


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/health')
def health():
    """Health check — useful for verifying the app started correctly."""
    return jsonify({
        'status': 'ok',
        'jobs_store': len(jobs),
        'databricks_configured': bool(DATABRICKS_HOST and DATABRICKS_TOKEN and DATABRICKS_WAREHOUSE_ID),
        'databricks_host': DATABRICKS_HOST[:30] + '...' if DATABRICKS_HOST else 'NOT SET',
        'groq_configured': bool(GROQ_API_KEY),
        'mcp_configured': bool(MCP_API_KEY),
    })


@app.route('/api/search-employee', methods=['POST'])
def search_employee():
    """Search for an employee by name."""
    name = request.json.get('name', '').strip()
    if not name:
        return jsonify({'error': 'Name is required'}), 400
    result = call_mcp_tool('search_employees', {'query': name})
    return jsonify(result)


@app.route('/api/aircall-users', methods=['GET'])
def get_aircall_users():
    """Get Aircall users list."""
    result = call_mcp_tool('list_aircall_users', {})
    return jsonify(result)


def fetch_calls_via_mcp(user_email: str, date_from: str) -> list:
    """Fallback: fetch calls via MCP API with date-range chunking."""
    from_date = datetime.strptime(date_from, '%Y-%m-%d')
    to_date = datetime.now()
    total_days = (to_date - from_date).days

    if total_days <= 3:
        result = call_mcp_tool('list_aircall_calls', {
            'user_email': user_email, 'limit': 100, 'date_from': date_from,
        })
        return _parse_aircall_calls(result)

    chunk_days = 5
    chunks = []
    current = from_date
    while current < to_date:
        chunk_end = min(current + timedelta(days=chunk_days), to_date)
        chunks.append((current.strftime('%Y-%m-%d'), chunk_end.strftime('%Y-%m-%d')))
        current = chunk_end

    all_calls = []
    seen_ids = set()

    def fetch_chunk(chunk_from, chunk_to):
        args = {'user_email': user_email, 'limit': 100, 'date_from': chunk_from}
        if chunk_to != to_date.strftime('%Y-%m-%d'):
            args['date_to'] = chunk_to
        return call_mcp_tool('list_aircall_calls', args)

    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = [pool.submit(fetch_chunk, cf, ct) for cf, ct in chunks]
        for future in futures:
            try:
                result = future.result()
                for c in _parse_aircall_calls(result):
                    cid = c.get('id') or c.get('call_id') or ''
                    if cid and cid not in seen_ids:
                        seen_ids.add(cid)
                        all_calls.append(c)
                    elif not cid:
                        all_calls.append(c)
            except Exception as e:
                logger.warning("Failed to fetch call chunk: %s", e)

    return all_calls


def fetch_eligible_calls(user_email: str, date_from: str) -> list:
    """Fetch calls pre-filtered against Zoho.
    Uses Databricks if available (has actual caller phone for filtering).
    Falls back to MCP API if Databricks is not configured.
    """
    use_databricks = all([DATABRICKS_HOST, DATABRICKS_TOKEN, DATABRICKS_WAREHOUSE_ID])

    if use_databricks:
        # --- Databricks path: full Zoho filtering ---
        logger.info("Fetching calls via Databricks for %s from %s", user_email, date_from)
        calls_rows = query_databricks(f"""
            SELECT
                c.id,
                c.raw_digits,
                c.direction,
                c.duration,
                CAST(c.started_at AS STRING) as started_at,
                CAST(c.answered_at AS STRING) as answered_at,
                CAST(c.ended_at AS STRING) as ended_at,
                c.recording,
                c.asset,
                c.contact.name as contact_name,
                c.status,
                CASE WHEN t.call_id IS NOT NULL THEN 'true' ELSE 'false' END as has_transcript
            FROM trilogycare_dev.aircall.calls c
            LEFT JOIN trilogycare_dev.aircall.transcriptions t ON c.id = t.call_id
            WHERE c.user.email = '{user_email}'
            AND c.duration >= 300
            AND c.started_at >= TIMESTAMP '{date_from} 00:00:00'
            ORDER BY c.started_at DESC
        """)

        logger.info("Databricks returned %d calls for %s", len(calls_rows), user_email)
        if not calls_rows:
            return []

        lead_phones = get_active_lead_phones()
        deal_phone_dates = get_deal_phones_by_date()

        eligible_calls = []
        for row in calls_rows:
            call_id = str(row[0]) if row[0] else ''
            raw_digits = row[1] or ''
            direction = row[2] or ''
            duration = row[3] or 0
            started_at = row[4] or ''
            answered_at = row[5] or ''
            ended_at = row[6] or ''
            recording = row[7] or ''
            asset = row[8] or ''
            contact_name = row[9] or ''
            status = row[10] or ''
            has_transcript = row[11] if len(row) > 11 else 'false'

            if not raw_digits or raw_digits == 'anonymous':
                continue

            caller_normalized = normalize_phone(raw_digits)
            if not caller_normalized:
                continue

            call_date = started_at[:10] if started_at else ''

            if caller_normalized not in lead_phones:
                if not (caller_normalized in deal_phone_dates and call_date in deal_phone_dates.get(caller_normalized, set())):
                    continue

            duration_int = int(duration) if duration else 0
            duration_mins = round(duration_int / 60, 1)
            fmt_started = started_at.replace(' ', 'T') + '+10:00' if started_at and '+' not in started_at and 'T' not in started_at else started_at
            fmt_answered = answered_at.replace(' ', 'T') + '+10:00' if answered_at and '+' not in answered_at and 'T' not in answered_at else answered_at
            fmt_ended = ended_at.replace(' ', 'T') + '+10:00' if ended_at and '+' not in ended_at and 'T' not in ended_at else ended_at
            call_obj = {
                'id': call_id,
                'call_id': call_id,
                'raw_digits': raw_digits,
                'phone_number': raw_digits,
                'direction': direction,
                'duration': duration_int,
                'duration_minutes': duration_mins,
                'created_at': fmt_started,
                'started_at': fmt_started,
                'answered_at': fmt_answered,
                'ended_at': fmt_ended,
                'status': status,
                'contact_name': contact_name,
                'recording': recording,
                'recording_url': recording or asset or '',
                'asset': asset or f"https://assets.aircall.io/calls/{call_id}/recording",
                'has_recording': bool(recording or asset),
                'has_transcript': has_transcript,
                'zoho_verified': True,
            }
            eligible_calls.append(call_obj)

        return eligible_calls

    else:
        # --- MCP fallback: fetch via API, show all 5+ min calls ---
        logger.info("Databricks not configured, fetching calls via MCP for %s", user_email)
        all_calls = fetch_calls_via_mcp(user_email, date_from)

        calls = []
        for c in all_calls:
            dur = c.get('duration') or 0
            if isinstance(dur, str):
                try:
                    dur = int(dur)
                except ValueError:
                    dur = 0
            if dur < 300:
                continue
            if c.get('recording') and not c.get('recording_url'):
                c['recording_url'] = c['recording']
            if not c.get('recording_url') and c.get('asset'):
                c['recording_url'] = c['asset']
            if c.get('recording') or c.get('recording_url') or c.get('asset'):
                c['has_recording'] = True
            calls.append(c)

        return calls


@app.route('/api/calls', methods=['POST'])
def get_calls():
    """Get recent calls for an Aircall user, pre-filtered against Zoho eligibility."""
    user_email = request.json.get('user_email', '').strip()
    date_filter = request.json.get('date', '').strip()
    if not user_email:
        return jsonify({'error': 'user_email is required'}), 400

    if date_filter == 'today' or not date_filter:
        date_from = datetime.now().strftime('%Y-%m-%d')
    elif date_filter == 'week':
        date_from = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
    else:
        date_from = (datetime.now() - timedelta(days=90)).strftime('%Y-%m-%d')

    calls = fetch_eligible_calls(user_email, date_from)

    today_str = datetime.now().strftime('%Y-%m-%d')
    unique_clients = set()
    today_calls = 0
    for c in calls:
        call_date = c.get('created_at', '')[:10]
        if call_date == today_str:
            today_calls += 1
            client_id = c.get('contact_name') or c.get('raw_digits') or ''
            if client_id:
                unique_clients.add(client_id.lower())

    return jsonify({
        'result': {
            'calls': calls,
            'stats': {
                'total_calls': len(calls),
                'today_calls': today_calls,
                'unique_clients_today': len(unique_clients),
                'date_from': date_from,
            }
        }
    })


@app.route('/api/start-review', methods=['POST'])
def start_review():
    """Kick off the review pipeline in a background thread. Returns a job_id for polling."""
    data = request.json
    job_id = str(uuid.uuid4())[:8]
    jobs[job_id] = {
        'status': 'running',
        'message': 'Fetching call transcript...',
        'result': None,
        'error': None,
    }

    thread = threading.Thread(target=run_review_pipeline, args=(job_id, data), daemon=True)
    thread.start()

    return jsonify({'job_id': job_id})


@app.route('/api/review-status/<job_id>', methods=['GET'])
def review_status(job_id):
    """Poll for the status of a review job."""
    job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    return jsonify(job)



@app.route('/api/start-review-from-transcript', methods=['POST'])
def start_review_from_transcript():
    """Start a review from an uploaded transcript file. Extracts rep name from content."""
    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({'error': 'Please upload a transcript file.'}), 400
    try:
        raw = file.read()
        transcript = raw.decode('utf-8', errors='replace').strip()
    except Exception as e:
        return jsonify({'error': f'Could not read file: {str(e)}'}), 400
    if len(transcript) < 100:
        return jsonify({'error': 'Transcript seems too short. Please upload the full call transcript.'}), 400

    rep_name = ''
    m = re.search(r'(?:agent|rep|representative|advisor)[:\s]+([A-Z][a-z]+ [A-Z][a-z]+)', transcript[:2000], re.IGNORECASE)
    if m: rep_name = m.group(1).strip()
    if not rep_name:
        m = re.search(r'\[agent\]\s*([A-Z][a-z]+ ?[A-Z]?[a-z]*)', transcript[:2000])
        if m: rep_name = m.group(1).strip()
    if not rep_name:
        m = re.search(r"(?:my name is|this is|I'm|I am)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s+(?:from|at|with)\s+Trilogy", transcript[:3000], re.IGNORECASE)
        if m: rep_name = m.group(1).strip()
    if not rep_name:
        base = os.path.splitext(file.filename)[0]
        parts = re.split(r'[_\-\s]+', base)
        name_parts = [p for p in parts if p[0:1].isupper() and len(p) > 1 and p.lower() not in ('call', 'transcript', 'review', 'recording', 'aircall')]
        if len(name_parts) >= 2: rep_name = ' '.join(name_parts[:2])
    if not rep_name:
        return jsonify({'error': 'Could not detect the rep name from the transcript. Please rename the file to include the rep name (e.g. "Sarah Palmer transcript.txt").'}), 400

    rep_email = ''
    rep_role = ''
    est_id = ''
    emp_result = call_mcp_tool('search_employees', {'query': rep_name})
    if isinstance(emp_result, dict) and not emp_result.get('error'):
        employees = []
        inner = emp_result.get('result', emp_result)
        if isinstance(inner, dict): employees = inner.get('employees', inner.get('results', []))
        elif isinstance(inner, list): employees = inner
        for emp in (employees if isinstance(employees, list) else []):
            emp_name = emp.get('full_name', '') or emp.get('name', '')
            if emp_name and rep_name.lower() in emp_name.lower():
                title_raw = emp.get('title', '')
                est_match = re.match(r'(EST-\d+)\s*(.*)', title_raw)
                est_id = emp.get('employee_id', '') or emp.get('est_id', '') or (est_match.group(1) if est_match else '')
                rep_role = emp.get('role', '') or emp.get('job_title', '') or (est_match.group(2) if est_match else title_raw)
                rep_email = emp.get('email', '') or emp.get('work_email', '')
                fn = emp.get('full_name', '') or emp.get('name', '')
                if fn: rep_name = fn
                break

    job_id = str(uuid.uuid4())[:8]
    jobs[job_id] = {'status': 'running', 'message': 'AI is scoring the call...', 'result': None, 'error': None}
    pipeline_data = {
        'rep_name': rep_name, 'rep_email': rep_email,
        'rep_role': rep_role, 'est_id': est_id,
        'call_date': datetime.now().strftime('%d %b %Y'),
        'call_direction': '', 'call_duration': '',
        'client_name': 'Unknown', 'client_phone': '',
        'transcript_text': transcript,
    }
    thread = threading.Thread(target=run_review_pipeline, args=(job_id, pipeline_data), daemon=True)
    thread.start()
    return jsonify({'job_id': job_id, 'rep_name': rep_name, 'rep_email': rep_email, 'rep_role': rep_role, 'est_id': est_id})


@app.route('/api/start-review-from-recording', methods=['POST'])
def start_review_from_recording():
    """Start a review from an uploaded audio recording. Transcribes via Groq Whisper then scores."""
    import requests as req_lib
    import tempfile

    if not GROQ_API_KEY:
        return jsonify({'error': 'Transcription service not configured (GROQ_API_KEY missing).'}), 500

    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({'error': 'Please upload a recording file (.mp3, .wav, .m4a).'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.mp3', '.wav', '.m4a', '.ogg', '.webm'):
        return jsonify({'error': f'Unsupported audio format: {ext}. Please upload .mp3, .wav, or .m4a.'}), 400

    tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
    file.save(tmp.name)
    tmp_path = tmp.name

    rep_name = ''
    base = os.path.splitext(file.filename)[0]
    parts = re.split(r'[_\-\s]+', base)
    name_parts = [p for p in parts if p[0:1].isupper() and len(p) > 1 and p.lower() not in ('call', 'transcript', 'review', 'recording', 'aircall', 'audio')]
    if len(name_parts) >= 2:
        rep_name = ' '.join(name_parts[:2])

    rep_email = ''
    rep_role = ''
    est_id = ''
    if rep_name:
        emp_result = call_mcp_tool('search_employees', {'query': rep_name})
        if isinstance(emp_result, dict) and not emp_result.get('error'):
            employees = []
            inner = emp_result.get('result', emp_result)
            if isinstance(inner, dict): employees = inner.get('employees', inner.get('results', []))
            elif isinstance(inner, list): employees = inner
            for emp in (employees if isinstance(employees, list) else []):
                emp_name = emp.get('full_name', '') or emp.get('name', '')
                if emp_name and rep_name.lower() in emp_name.lower():
                    title_raw = emp.get('title', '')
                    est_match = re.match(r'(EST-\d+)\s*(.*)', title_raw)
                    est_id = emp.get('employee_id', '') or emp.get('est_id', '') or (est_match.group(1) if est_match else '')
                    rep_role = emp.get('role', '') or emp.get('job_title', '') or (est_match.group(2) if est_match else title_raw)
                    rep_email = emp.get('email', '') or emp.get('work_email', '')
                    fn = emp.get('full_name', '') or emp.get('name', '')
                    if fn: rep_name = fn
                    break

    job_id = str(uuid.uuid4())[:8]
    jobs[job_id] = {'status': 'running', 'message': 'Transcribing recording...', 'result': None, 'error': None}

    def transcribe_and_review():
        try:
            mime_types = {'.mp3': 'audio/mpeg', '.wav': 'audio/wav', '.m4a': 'audio/mp4', '.ogg': 'audio/ogg', '.webm': 'audio/webm'}
            mime = mime_types.get(ext, 'audio/mpeg')
            with open(tmp_path, 'rb') as audio_file:
                result = req_lib.post(
                    "https://api.groq.com/openai/v1/audio/transcriptions",
                    headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                    files={"file": (file.filename, audio_file, mime)},
                    data={"model": "whisper-large-v3", "language": "en", "response_format": "text"},
                    timeout=300,
                )
                result.raise_for_status()
            os.unlink(tmp_path)
            transcript = result.text.strip()
            if not transcript:
                jobs[job_id] = {'status': 'error', 'message': 'Transcription returned empty. The recording may be silent or corrupted.', 'result': None, 'error': 'Empty transcription'}
                return
            jobs[job_id]['message'] = 'Transcription complete. AI is scoring the call...'
            pipeline_data = {
                'rep_name': rep_name or 'Unknown Rep', 'rep_email': rep_email,
                'rep_role': rep_role, 'est_id': est_id,
                'call_date': datetime.now().strftime('%d %b %Y'),
                'call_direction': '', 'call_duration': '',
                'client_name': 'Unknown', 'client_phone': '',
                'transcript_text': transcript,
            }
            run_review_pipeline(job_id, pipeline_data)
        except Exception as e:
            logger.error('Recording transcription failed: %s', e)
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            jobs[job_id] = {'status': 'error', 'message': f'Transcription failed: {str(e)}', 'result': None, 'error': str(e)}

    thread = threading.Thread(target=transcribe_and_review, daemon=True)
    thread.start()
    return jsonify({'job_id': job_id, 'rep_name': rep_name, 'rep_email': rep_email, 'rep_role': rep_role, 'est_id': est_id})


def run_review_pipeline(job_id: str, data: dict):
    """Background worker: fetch transcript, score with LLM, generate doc."""
    try:
        call_id = data.get('call_id')
        rep_name = data.get('rep_name', '')
        rep_email = data.get('rep_email', '')
        rep_role = data.get('rep_role', '')
        est_id = data.get('est_id', '')
        call_date = data.get('call_date', '')
        call_direction = data.get('call_direction', '')
        call_duration = data.get('call_duration', '')
        client_name = data.get('client_name', 'Unknown')
        client_phone = data.get('client_phone', '')
        direct_link = data.get('direct_link', '')
        recording_url = data.get('recording_url', '')

        # Zoho eligibility is already verified at the call-list stage (Databricks filtering).
        # If a client phone is available, try to resolve the client name from Zoho.
        if client_phone and client_name == 'Unknown':
            normalized = normalize_phone(client_phone)
            if normalized:
                rows = query_databricks(f"""
                    SELECT First_Name, Last_Name
                    FROM trilogycare_dev.zoho_crm.leads
                    WHERE Journey_Stage IN ('B-Active HCP', 'B-Allocated HCP')
                    AND (
                        RIGHT(REGEXP_REPLACE(Phone, '[^0-9]', ''), 9) = '{normalized}'
                        OR RIGHT(REGEXP_REPLACE(Mobile, '[^0-9]', ''), 9) = '{normalized}'
                        OR RIGHT(REGEXP_REPLACE(Contact_Phone, '[^0-9]', ''), 9) = '{normalized}'
                    )
                    LIMIT 1
                """)
                if rows:
                    client_name = f"{rows[0][0] or ''} {rows[0][1] or ''}".strip() or client_name

        # Step 1: Get transcript (pre-provided or fetched from Aircall)
        transcript_text = data.get('transcript_text', '')
        resolved_client = client_name

        def fetch_transcript():
            if data.get('transcript_text'):
                return data['transcript_text']
            # Priority 1: Query Databricks transcriptions table directly
            if call_id:
                db_transcript = get_transcript_from_databricks(str(call_id))
                if db_transcript:
                    return db_transcript
            # Priority 2: get_aircall_call_details (has segments for processed calls)
            call_data = get_call_details(str(call_id))
            if call_data:
                transcript = call_data.get('transcript', {})
                segments = transcript.get('segments', []) if isinstance(transcript, dict) else []
                if segments:
                    return '\n'.join(
                        f"[{seg.get('participant_type', 'unknown')}]: {seg.get('text', '')}"
                        for seg in segments if seg.get('text')
                    )
            # Priority 3: get_aircall_transcript MCP tool
            result = call_mcp_tool('get_aircall_transcript', {'call_id': str(call_id)})
            if isinstance(result, dict):
                inner = result.get('result', result)
                if isinstance(inner, dict):
                    segs = inner.get('transcript_segments', [])
                    if segs:
                        return '\n'.join(
                            f"[{seg.get('participant_type', 'unknown')}]: {seg.get('text', '')}"
                            for seg in segs if seg.get('text')
                        )
                    return inner.get('raw', '') or inner.get('transcript', '') or inner.get('text', '') or ''
            return ''

        def fetch_client():
            if client_name != 'Unknown' or not client_phone:
                return None
            return call_mcp_tool('search_customers', {'query': client_phone})

        with ThreadPoolExecutor(max_workers=2) as pool:
            future_transcript = pool.submit(fetch_transcript)
            future_client = pool.submit(fetch_client)

            transcript_text = future_transcript.result() or ''

            # If no stored transcript, transcribe from recording (live via Groq Whisper)
            if not transcript_text and (recording_url or call_id):
                jobs[job_id]['message'] = 'No stored transcript — transcribing from recording...'
                logger.info('No transcript in API for call %s, attempting live transcription (recording_url=%s, call_id=%s)',
                            call_id, bool(recording_url), call_id)
                transcript_text = transcribe_recording(recording_url, call_id=str(call_id) if call_id else '')

            if not transcript_text:
                if not GROQ_API_KEY:
                    msg = 'Transcription service not configured (GROQ_API_KEY missing). Transcripts are batched at 6am — try again tomorrow or upload a transcript file.'
                else:
                    msg = 'Could not get transcript for this call. The recording may still be processing — try again in a few minutes, or select an older call.'
                logger.error('No transcript available for call %s', call_id)
                jobs[job_id] = {'status': 'error', 'message': msg, 'result': None, 'error': 'No transcript'}
                return

            client_result = future_client.result()
            if client_result and isinstance(client_result, dict) and not client_result.get('error'):
                raw = client_result.get('raw', '') or json.dumps(client_result)
                resolved_client = extract_client_name(raw) or client_name

        # Step 2: Score with LLM
        jobs[job_id]['message'] = 'AI is scoring the call across 4 stages...'

        scoring_prompt = f"""Score this call quality review.

REP (the Trilogy Care employee being reviewed): {rep_name} ({rep_role})
CLIENT (the external caller/recipient): {resolved_client}
CALL DATE: {call_date}
DIRECTION: {call_direction}
DURATION: {call_duration} minutes

IMPORTANT: In the transcript below, [agent] is ALWAYS {rep_name} (the rep being reviewed). [external] is ALWAYS the client/caller. Do NOT confuse these roles. When writing your review, refer to {rep_name} as the rep and {resolved_client} as the client.

TRANSCRIPT:
{transcript_text[:12000]}

Score each dimension 1-10 based on the rubric. Be specific and reference real moments from the transcript. Return valid JSON only."""

        try:
            llm_response = call_llm(SCORING_SYSTEM_PROMPT, scoring_prompt, max_tokens=6000)
            logger.info("LLM response length: %d chars, first 200: %s", len(llm_response), llm_response[:200])
            review_json = parse_json_from_response(llm_response)
        except Exception as e:
            logger.error("LLM scoring exception: %s", e)
            jobs[job_id] = {'status': 'error', 'message': f"LLM scoring failed: {str(e)}", 'result': None, 'error': str(e)}
            return

        if not review_json or not review_json.get('scores'):
            logger.error("LLM returned invalid JSON. Raw response (first 500): %s", llm_response[:500] if llm_response else 'EMPTY')
            jobs[job_id] = {'status': 'error', 'message': 'LLM returned invalid or incomplete JSON. Please retry.', 'result': None, 'error': 'Invalid JSON from LLM'}
            return

        # Use LLM-identified client name if still Unknown
        llm_client = review_json.get('client_name', '')
        if (not resolved_client or resolved_client == 'Unknown') and llm_client:
            resolved_client = llm_client

        # Step 3: Generate Word doc
        jobs[job_id]['message'] = 'Generating Word document...'

        review_data = {
            'rep_name': rep_name,
            'rep_email': rep_email,
            'rep_role': rep_role,
            'est_id': est_id,
            'call_id': call_id,
            'call_date': call_date,
            'direction': call_direction,
            'duration': call_duration,
            'client_name': resolved_client,
            'client_phone': client_phone,
            'scores': review_json.get('scores', {}),
            'stage_flow': review_json.get('stage_flow', ''),
            'focus_on': review_json.get('focus_on', ''),
            'top_strength': review_json.get('top_strength', ''),
            'top_development_area': review_json.get('top_development_area', ''),
        }

        try:
            filepath = generate_word_doc(review_data)
        except Exception as e:
            jobs[job_id] = {'status': 'error', 'message': f"Document generation failed: {str(e)}", 'result': None, 'error': str(e)}
            return

        call_link = f"https://assets.aircall.io/calls/{call_id}/recording"
        review_data['call_link'] = call_link
        review_data['filepath'] = filepath
        review_data['filename'] = os.path.basename(filepath)

        # Document ready — attachment will be sent inline via MS Graph at email time

        jobs[job_id] = {'status': 'complete', 'message': 'Review complete!', 'result': review_data, 'error': None}

    except Exception as e:
        jobs[job_id] = {'status': 'error', 'message': f"Unexpected error: {str(e)}", 'result': None, 'error': str(e)}


@app.route('/api/send-email', methods=['POST'])
def send_email_route():
    """Send the review document via email with .docx attachment using Microsoft Graph."""
    import requests as req_lib
    data = request.json
    rep_email = data.get('rep_email', '')
    rep_name = data.get('rep_name', '')
    first_name = rep_name.split()[0] if rep_name else ''
    client_name = data.get('client_name', '')
    call_date = data.get('call_date', '')
    call_link = data.get('call_link', '')
    filepath = data.get('filepath', '')
    filename = data.get('filename', '')

    if not rep_email:
        return jsonify({'error': 'No recipient email address provided.'}), 400

    subject = f"Call Quality Review \u2014 {client_name} \u2014 {call_date}"
    body_text = f"""Hi {first_name},

Please find attached your call quality review for your call with {client_name} on {call_date}:
{call_link}

Have a read through and let me know if you'd like to chat about anything in the review.

Kind regards,
Anesu"""

    # Build Microsoft Graph sendMail payload with inline attachment
    mail_payload = {
        'message': {
            'subject': subject,
            'body': {'contentType': 'Text', 'content': body_text},
            'from': {'emailAddress': {'address': SENDER_EMAIL}},
            'toRecipients': [{'emailAddress': {'address': rep_email}}],
        },
        'saveToSentItems': True,
    }

    # Attach the Word document if available on disk
    attachment_ok = False
    if filepath and os.path.isfile(filepath):
        try:
            with open(filepath, 'rb') as f:
                file_bytes = f.read()
            mail_payload['message']['attachments'] = [{
                '@odata.type': '#microsoft.graph.fileAttachment',
                'name': filename or os.path.basename(filepath),
                'contentBytes': base64.b64encode(file_bytes).decode('utf-8'),
                'contentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            }]
            attachment_ok = True
            logger.info('Attached %s (%d bytes)', filename, len(file_bytes))
        except Exception as e:
            logger.error('Failed to read file for attachment: %s', e)

    try:
        # Get Azure AD token
        token = get_graph_token()
        if not token:
            return jsonify({'error': 'Failed to get Microsoft Graph token'}), 500

        resp = req_lib.post(
            f'https://graph.microsoft.com/v1.0/users/{SENDER_EMAIL}/sendMail',
            json=mail_payload,
            headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
            timeout=30,
        )
        logger.info('Graph sendMail status: %s', resp.status_code)
        if resp.status_code == 202:
            return jsonify({'success': True, 'attachment_included': attachment_ok})
        else:
            error_text = resp.text[:500]
            logger.error('Graph sendMail failed: %s %s', resp.status_code, error_text)
            return jsonify({'error': f'Microsoft Graph error: {resp.status_code}'}), 500
    except Exception as e:
        logger.error('send_email failed: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/output/<filename>')
def download_file(filename):
    """Download a generated document."""
    return send_from_directory(app.config['OUTPUT_DIR'], filename, as_attachment=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_graph_token():
    """Get a Microsoft Graph API access token via Azure AD client credentials."""
    import requests as req_lib
    if not all([AZURE_AD_TENANT_ID, AZURE_AD_CLIENT_ID, AZURE_AD_CLIENT_SECRET]):
        logger.error('Azure AD credentials not configured')
        return None
    try:
        resp = req_lib.post(
            f'https://login.microsoftonline.com/{AZURE_AD_TENANT_ID}/oauth2/v2.0/token',
            data={
                'grant_type': 'client_credentials',
                'client_id': AZURE_AD_CLIENT_ID,
                'client_secret': AZURE_AD_CLIENT_SECRET,
                'scope': 'https://graph.microsoft.com/.default',
            },
            timeout=10,
        )
        data = resp.json()
        return data.get('access_token')
    except Exception as e:
        logger.error('Failed to get Graph token: %s', e)
        return None


def extract_client_name(raw: str) -> str:
    """Try to extract a client name from search results (various MCP response shapes)."""
    try:
        data = json.loads(raw)
        # Unwrap nested MCP response: {result: {customers: [...]}} or {result: [...]}
        if isinstance(data, dict):
            inner = data.get('result', data)
            if isinstance(inner, dict):
                # {result: {customers: [...]}} or {customers: [...]}
                for key in ('customers', 'contacts', 'results', 'data'):
                    if isinstance(inner.get(key), list) and inner[key]:
                        inner = inner[key]
                        break
            if isinstance(inner, list) and inner:
                item = inner[0]
                return (item.get('display_name', '') or item.get('full_name', '')
                        or item.get('name', '')
                        or f"{item.get('first_name', '')} {item.get('last_name', '')}".strip())
            if isinstance(inner, dict):
                return (inner.get('display_name', '') or inner.get('full_name', '')
                        or inner.get('name', '')
                        or f"{inner.get('first_name', '')} {inner.get('last_name', '')}".strip())
        if isinstance(data, list) and data:
            item = data[0]
            return (item.get('display_name', '') or item.get('full_name', '')
                    or item.get('name', '')
                    or f"{item.get('first_name', '')} {item.get('last_name', '')}".strip())
    except Exception:
        pass
    return ''


def parse_json_from_response(text: str) -> dict:
    """Extract JSON from LLM response (may be wrapped in markdown code blocks)."""
    try:
        return json.loads(text)
    except Exception:
        pass
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            pass
    return {}


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5050))
    debug = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug)
